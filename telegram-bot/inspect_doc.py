"""讀取合約 docx，印出所有段落內容供分析。"""
from pathlib import Path
from docx import Document

docx_path = Path.home() / "Desktop" / "克勞德" / "2026日日好日續約工商登記合約.docx"

if not docx_path.exists():
    print(f"找不到檔案：{docx_path}")
    exit(1)

doc = Document(docx_path)

print("=== 段落內容 ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"[{i}] {repr(para.text)}")

print("\n=== 表格內容 ===")
for t_idx, table in enumerate(doc.tables):
    print(f"\n表格 {t_idx}：")
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text:
                print(f"  [{r_idx},{c_idx}] {repr(text)}")
